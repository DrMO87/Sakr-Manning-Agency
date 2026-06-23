import docx
from docx.document import Document
from datetime import date
from .schemas import ContractData

def fill_contract(template_path: str, output_path: str, data: ContractData):
    doc = docx.Document(template_path)
    
    # We found that Table 2 (index 2) contains most sections
    if len(doc.tables) < 3:
        raise ValueError("Document does not have enough tables (expected at least 3)")
    
    table = doc.tables[2]
    
    # Helper to find row by first cell text (contains)
    def find_row_index(start_row, keyword):
        for i in range(start_row, len(table.rows)):
            # Check all cells in row for keyword
            for cell in table.rows[i].cells:
                if keyword.lower() in cell.text.lower():
                    return i
        return -1

    # Simple helper to set text in next empty cell
    def set_value_after_label(row_idx, label_keyword, value):
        if row_idx == -1: return
        row = table.rows[row_idx]
        found_label = False
        for i, cell in enumerate(row.cells):
            if label_keyword.lower() in cell.text.lower():
                found_label = True
                continue # Move to next cell
            
            if found_label:
                # We are in cells after the label. Find the first one that is "empty" or just the immediate next?
                # Merged cells often repeat the SAME cell object in python-docx list.
                # We need to find a cell that is NOT the same object as the label cell? 
                # Or just the next index.
                # Let's try inserting into the next index that isn't the label.
                try:
                    cell.text = str(value) if value is not None else ""
                    return # Done
                except Exception:
                    continue

    # --- 1. PERSONAL DETAILS ---
    # Find "Full Name" row
    fullname_row = find_row_index(0, "Full Name")
    # Note: "Full Name" might appear multiple times (Next of Kin). We want the first one.
    if fullname_row != -1:
        # Assuming layout: Full Name | Value | Date Of Birth | Value
        # But cell indices might be tricky. Let's try direct setting if we know the offset.
        # Based on inspection, we might need custom offsets.
        # Let's use a simpler strategy: Scan cells, finding label, then set NEXT cell.
        pass

    # Strategies:
    # Row "Full Name": Label(0), Value(1), Label(2), Value(3) ?
    # Let's try to set specific values based on our variable search logic
    
    def set_next_cell(row_idx, label, value, skip=0):
        if row_idx == -1 or value is None: return
        row = table.rows[row_idx]
        for i, cell in enumerate(row.cells):
            if label.lower() in cell.text.lower():
                # Target is i + 1 + skip
                target_idx = i + 1 + skip
                if target_idx < len(row.cells):
                    row.cells[target_idx].text = str(value)
                    return

    # Personal
    p_row = find_row_index(0, "Full Name")
    set_next_cell(p_row, "Full Name", data.full_name) # Name
    set_next_cell(p_row, "Date Of Birth", data.date_of_birth) # DOB
    
    nat_row = find_row_index(0, "Nationality")
    set_next_cell(nat_row, "Nationality", data.nationality)
    set_next_cell(nat_row, "Height", data.height_cm, skip=5) # Based on 'Cell 7' vs 'Cell 1' gap seen in logs? 
    # Actually, let's just try adjacent first. If merged, next index is the merged cell.
    # The logs showed "Height" at 7. It implies huge merges.
    # Let's try safer approach: set_value_after_label
    
    def safe_set(row_idx, label, value):
        if row_idx == -1: return
        row = table.rows[row_idx]
        found = False
        for cell in row.cells:
            if label.lower() in cell.text.lower():
                found = True
            elif found and not cell.text.strip(): # Empty cell after label
                cell.text = str(value)
                return

    # Re-apply with safe_set
    p_row = find_row_index(0, "Full Name")
    if p_row != -1:
         # Manually handling Row X mapping based on observation
         # If we can't find exact empty cell, we might overwrite specific indices
         # Let's try safe_set which looks for empty cell
         safe_set(p_row, "Full Name", data.full_name)
         safe_set(p_row, "Date Of Birth", data.date_of_birth)

    nat_row = find_row_index(0, "Nationality")
    safe_set(nat_row, "Nationality", data.nationality)
    safe_set(nat_row, "Height", data.height_cm)
    safe_set(nat_row, "Weight", data.height_cm)
    
    place_row = find_row_index(0, "Place Of Birth")
    safe_set(place_row, "Place Of Birth", data.place_of_birth)
    
    shirt_row = find_row_index(0, "Shirt Size")
    safe_set(shirt_row, "Shirt Size", data.shirt_size)
    
    trouser_row = find_row_index(0, "Trouser Size")
    safe_set(trouser_row, "Trouser Size", data.trouser_size)
    safe_set(trouser_row, "Shoes Size", data.shoe_size)

    # --- 2. EDUCATION ---
    edu_row = find_row_index(5, "College / School")
    safe_set(edu_row, "College", data.education)
    
    english_row = find_row_index(8, "English Language")
    # Checkbox logic handling? Or just write "X"
    # "Fluent [] Good []" -> we need to find the cell with text "Good" and maybe append "X" or separate cell?
    # Skipping checkboxes for now, assumed text entry
    
    # --- 3. CONTACT ---
    addr_row = find_row_index(10, "Home Address")
    safe_set(addr_row, "Address", data.address)
    
    email_row = find_row_index(12, "E-Mail")
    safe_set(email_row, "E-Mail", data.email)
    safe_set(email_row, "Mobile", data.phone_number)

    # --- 4. TRAVEL ---
    # Passport
    pass_row = find_row_index(15, "Passport")
    if pass_row != -1 and data.passport_number:
        # Assuming Passport row has columns for No, Issue, Expire
        # Index guessing: [Label, No, Issue, Expire, By, Place]
        # Let's just blindly fill the next few empty cells
        cells = table.rows[pass_row].cells
        filled_count = 0
        for cell in cells:
            if "Passport" in cell.text: continue
            if not cell.text.strip():
                if filled_count == 0: cell.text = str(data.passport_number)
                elif filled_count == 1: cell.text = str(data.passport_issue_date)
                elif filled_count == 2: cell.text = str(data.passport_expiry_date)
                elif filled_count == 3: cell.text = str(data.passport_issued_by)
                filled_count += 1
                if filled_count >= 4: break

    # Seaman Book
    sea_row = find_row_index(15, "Seaman Book")
    if sea_row != -1 and sea_row != pass_row and data.seaman_book_number:
         cells = table.rows[sea_row].cells
         filled_count = 0
         for cell in cells:
            if "Seaman Book" in cell.text: continue
            if not cell.text.strip() and "EAMS" not in cell.text: # Avoid overwriting 'EAMS' if prefilled
                if filled_count == 0: cell.text = str(data.seaman_book_number)
                elif filled_count == 1: cell.text = str(data.seaman_book_issue_date)
                elif filled_count == 2: cell.text = str(data.seaman_book_expiry_date)
                filled_count += 1
                if filled_count >= 3: break

    # --- 6. NOK ---
    nok_start = find_row_index(20, "NEXT OF KIN")
    if nok_start != -1:
        # Find Full Name after NOK header
        nok_name_row = find_row_index(nok_start, "Full Name")
        safe_set(nok_name_row, "Full Name", data.next_of_kin_name)
        safe_set(nok_name_row, "Relationship", data.next_of_kin_relationship)
        
        nok_addr_row = find_row_index(nok_name_row, "Address")
        safe_set(nok_addr_row, "Address", data.next_of_kin_address)
        
        nok_tel_row = find_row_index(nok_addr_row, "Tel")
        safe_set(nok_tel_row, "Tel", data.next_of_kin_phone)

    doc.save(output_path)
    return output_path

def _set_cell(table, row_idx, col_idx, value):
    # Deprecated
    pass

