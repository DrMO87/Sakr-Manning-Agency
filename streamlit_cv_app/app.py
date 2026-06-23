import streamlit as st
import pandas as pd
import json
import time
import os
import re
import requests
import google.generativeai as genai
import PyPDF2
import docx
from io import BytesIO

# --- AI Extraction Logic ---
def extract_text_from_pdf(pdf_file):
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    text = ""
    for page in pdf_reader.pages:
        extracted = page.extract_text()
        if extracted:
            text += extracted + "\n"
    return text

def extract_text_from_docx(docx_file):
    """Extract text from a DOCX file using python-docx."""
    try:
        doc = docx.Document(docx_file)
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
        return text
    except Exception as e:
        print(f"Warning: Failed to extract text from DOCX using python-docx: {e}")
        return f"[Error: Could not extract text from DOCX file. File may be corrupted or contain invalid media: {e}]"

def parse_cv_with_ai(file_bytes, cv_text, api_key, filename="", is_docx=False):
    genai.configure(api_key=api_key)
    model_name = None
    try:
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                if 'flash' in m.name or 'pro' in m.name:
                    model_name = m.name
                    break
    except Exception as e:
        return {"error": f"API Error (Likely invalid or expired API Key): {str(e)}"}
                
    if not model_name:
        return {"error": "No supported text generation models found for this API key."}
        
    model = genai.GenerativeModel(model_name)
    
    prompt = f"""
    You are an AI assistant that extracts ALL structured information from seafarer resumes/CVs.
    
    Filename: {filename}
    
    Extract the following details from the {'document text below' if is_docx else 'attached PDF document'}.
    
    You must return a JSON object with EXACTLY this structure (use null for unknown values, except where arrays/objects are specified):
    {{
        "user_name": "",
        "user_email_display": "",
        "position_name": "",
        "status": "Pending",
        "notes": "Auto-submitted via AI CV Extractor",
        "coded_rank": [
            {{
                "assigned_code": "",
                "rank_code": "",
                "rank_name": ""
            }}
        ],
        "user_documents": {{
            "passport": {{ "passport_no": null, "issue_date": null, "expiry_date": null, "issued_by": null, "place_of_issue": null }},
            "seaman_book": {{ "seaman_book_no": null, "issue_date": null, "expiry_date": null, "issued_by": null, "place_of_issue": null }},
            "coc": {{ "certificate_name": null, "certificate_number": null, "issue_date": null, "expiry_date": null, "issued_by": null, "issued_at": null }},
            "goc": {{ "certificate_number": null, "issue_date": null, "expiry_date": null, "issued_by": null, "issued_at": null }},
            "health_certificate": {{ "number": null, "issue_date": null, "expiry_date": null, "international_medical_number": null, "international_medical_issue_date": null, "international_medical_expiry_date": null }},
            "sea_service": [
                {{ "vessel_name": "", "rank": "", "vessel_type": "", "signed_on": "YYYY-MM-DD", "signed_off": "YYYY-MM-DD" }}
            ],
            "marine_courses": []
        }},
        "seafarer_application": {{
            "1_personal_details": {{
                "full_name": "",
                "date_of_birth": "YYYY-MM-DD",
                "marital_status": {{ "single": false, "married": false }},
                "nationality": "",
                "place_of_birth": ""
            }},
            "3_contact_details": {{
                "home_address_city": "",
                "e_mail": "",
                "mobile_tel": ""
            }},
            "8_marine_courses": [],
            "9_complete_sea_service_details": {{
                "service_records": [
                    {{
                        "company_name": "",
                        "rank": "",
                        "vessel_name_imo_number": "",
                        "signed_on": "YYYY-MM-DD",
                        "signed_off": "YYYY-MM-DD",
                        "period": "",
                        "vessel_type": "",
                        "dwt_grt": ""
                    }}
                ]
            }}
        }}
    }}
    
    CRITICAL RULES:
    - Use YYYY-MM-DD format for ALL dates. If a date cannot be formatted, return null.
    - Extract as many sea service records and marine courses as you can find.
    - If a field is not found in the CV, output null (for strings/numbers) or an empty string.
    - DO NOT include IDs (like "id": 110, or "user": 160) because those are database generated.
    """
    
    try:
        if is_docx:
            docx_prompt = prompt + f"\n\n--- DOCUMENT TEXT ---\n{cv_text}\n--- END OF DOCUMENT ---"
            response = model.generate_content(docx_prompt)
        else:
            response = model.generate_content([
                {"mime_type": "application/pdf", "data": file_bytes},
                prompt
            ])
        text_resp = response.text.strip()
        
        if text_resp.startswith("```json"):
            text_resp = text_resp[7:]
        elif text_resp.startswith("```"):
            text_resp = text_resp[3:]
            
        if text_resp.endswith("```"):
            text_resp = text_resp[:-3]
            
        return json.loads(text_resp.strip())
    except Exception as e:
        return {"error": str(e)}

# --- Streamlit UI ---
st.set_page_config(page_title="CV AI to JSON Converter", page_icon="🤖", layout="wide")

st.title("🤖 AI-Powered CV to JSON Extractor")
st.markdown("Upload CVs (PDF/DOCX) → Extract details into massive JSON → Save directly to `all_json_files` folder.")

# Sidebar for Configuration
with st.sidebar:
    st.header("⚙️ Configuration")
    api_key = st.text_input("Gemini API Key", type="password", help="Required for AI extraction.")

# Upload Area
st.subheader("📤 Upload Candidate CVs")

uploaded_files = st.file_uploader("Drop PDF or DOCX files here", type=["pdf", "docx"], accept_multiple_files=True)

if uploaded_files:
    if not api_key:
        st.warning("⚠️ Please enter your Gemini API Key in the sidebar.")
    else:
        if st.button("Extract & Save JSON", type="primary"):
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            # Create target directory
            project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            target_dir = os.path.join(project_root, "all_json_files")
            os.makedirs(target_dir, exist_ok=True)
            
            successful_uploads = []
            
            for i, file in enumerate(uploaded_files):
                status_text.text(f"Processing {file.name} ({i+1}/{len(uploaded_files)})...")
                
                # Step A: Read raw bytes and extract text
                file.seek(0)
                file_bytes = file.read()
                file.seek(0)
                
                file_ext = os.path.splitext(file.name)[1].lower()
                is_docx = file_ext == '.docx'
                
                if is_docx:
                    cv_text = extract_text_from_docx(file)
                else:
                    cv_text = extract_text_from_pdf(file)
                
                # Step B: AI Parse — extract ALL details
                result = parse_cv_with_ai(file_bytes, cv_text, api_key, filename=file.name, is_docx=is_docx)
                
                if "error" in result:
                    st.error(f"❌ Failed to parse {file.name}: {result['error']}")
                else:
                    # Save JSON to disk
                    base_name = os.path.splitext(file.name)[0]
                    output_file_name = f"{base_name}.json"
                    output_path = os.path.join(target_dir, output_file_name)
                    
                    with open(output_path, 'w', encoding='utf-8') as f:
                        json.dump(result, f, indent=4)
                    
                    st.success(f"✅ Saved to: {output_path}")
                    successful_uploads.append(output_file_name)
                    
                    with st.expander(f"🔍 AI Extraction for {file.name}", expanded=False):
                        st.json(result)
                
                progress_bar.progress((i + 1) / len(uploaded_files))
                
                if i < len(uploaded_files) - 1:
                    status_text.text(f"Pausing 15s for Gemini API limits (File {i+1} of {len(uploaded_files)} done)...")
                    time.sleep(15)
            
            status_text.text("✅ All processing complete!")
            if successful_uploads:
                st.info(f"Successfully saved {len(successful_uploads)} JSON files into {target_dir}")
