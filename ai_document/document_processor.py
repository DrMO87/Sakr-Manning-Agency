"""
Document processing utilities for extracting text and metadata from DOCX and PDF files.
Supports both synchronous and asynchronous processing.

Redesigned for 100% performance:
- Per-page structured extraction (text + tables per page)
- File size guard (20 MB limit)
- Improved OCR fallback threshold
- Fixed pdfplumber table/page association
- Fixed DOCX merged-cell handling
"""

import os
import logging
import hashlib
import zipfile
import tempfile
import shutil
from typing import Dict, List, Optional, Tuple
from pathlib import Path

# PDF processing
try:
    import fitz  # PyMuPDF - more robust PDF processing
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

# DOCX processing
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Alternative PDF processing
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False


logger = logging.getLogger(__name__)

# Maximum file size allowed for processing (20 MB)
MAX_FILE_SIZE_BYTES = 20 * 1024 * 1024

# Minimum word count to skip OCR fallback
MIN_WORDS_FOR_TEXT_PDF = 30


class DocumentProcessingError(Exception):
    """Custom exception for document processing errors."""
    pass


class DocumentProcessor:
    """
    Main class for processing documents (PDF and DOCX).
    Extracts text content and metadata from uploaded files.

    Returns structured per-page data to allow the extraction engine
    to send only the relevant sections to the LLM.
    """

    def __init__(self):
        self.supported_formats = []
        if PYMUPDF_AVAILABLE or PYPDF2_AVAILABLE or PDFPLUMBER_AVAILABLE:
            self.supported_formats.append('pdf')
        if DOCX_AVAILABLE:
            self.supported_formats.append('docx')

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def process_document(self, file_path: str, document_type: str = None) -> Dict:
        """
        Process a document file and extract text content and metadata.

        Args:
            file_path: Path to the document file.
            document_type: Optional override ('pdf' or 'docx').

        Returns:
            Dict with keys:
              - extracted_text (str): Full concatenated text from all pages.
              - page_count (int | None): Number of pages.
              - word_count (int): Word count of extracted_text.
              - pages (list): Per-page dicts: {"page": N, "text": "...", "tables": [[...]]}.
              - tables (list): All tables across all pages (2-D list of lists).
              - processing_error (str | None): Error message if any non-fatal issues.

        Raises:
            DocumentProcessingError: If the file cannot be processed at all.
        """
        if not os.path.exists(file_path):
            raise DocumentProcessingError(f"File not found: {file_path}")

        # File size guard
        file_size = os.path.getsize(file_path)
        if file_size > MAX_FILE_SIZE_BYTES:
            raise DocumentProcessingError(
                f"File too large: {file_size / 1024 / 1024:.1f} MB. "
                f"Maximum allowed size is {MAX_FILE_SIZE_BYTES // 1024 // 1024} MB."
            )

        if not document_type:
            document_type = self._detect_document_type(file_path)

        if document_type not in self.supported_formats:
            raise DocumentProcessingError(
                f"Unsupported document type: {document_type}. "
                f"Supported formats: {', '.join(self.supported_formats)}"
            )

        try:
            if document_type == 'pdf':
                return self._process_pdf(file_path)
            elif document_type == 'docx':
                return self._process_docx(file_path)
            else:
                raise DocumentProcessingError(f"Unknown document type: {document_type}")

        except DocumentProcessingError:
            raise
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            raise DocumentProcessingError(f"Failed to process document: {str(e)}")

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    def _detect_document_type(self, file_path: str) -> str:
        extension = Path(file_path).suffix.lower()
        if extension == '.pdf':
            return 'pdf'
        elif extension in ('.docx', '.doc'):
            return 'docx'
        else:
            raise DocumentProcessingError(f"Unsupported file extension: {extension}")

    # ------------------------------------------------------------------
    # PDF processing
    # ------------------------------------------------------------------

    def _process_pdf(self, file_path: str) -> Dict:
        """
        Process PDF file.

        Strategy:
          1. PyMuPDF for text (per page).
          2. pdfplumber for table extraction (per page) — tables are
             de-duplicated by content hash before being returned.
          3. If total extracted word count < MIN_WORDS_FOR_TEXT_PDF,
             fall back to Gemini OCR.
        """
        pages_data: List[Dict] = []          # per-page results
        all_tables: List[List[List[str]]] = []  # de-duplicated global table list
        seen_table_hashes = set()

        # ── Phase 1: Text via PyMuPDF ──────────────────────────────────
        if PYMUPDF_AVAILABLE:
            try:
                doc = fitz.open(file_path)
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    text = page.get_text()
                    pages_data.append({
                        "page": page_num + 1,
                        "text": text.strip(),
                        "tables": [],
                    })
                doc.close()
            except Exception as e:
                logger.warning(f"PyMuPDF failed: {e}")
        
        # Fallback: PyPDF2
        if not pages_data and PYPDF2_AVAILABLE:
            try:
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    for page_num, page in enumerate(reader.pages):
                        text = page.extract_text() or ""
                        pages_data.append({
                            "page": page_num + 1,
                            "text": text.strip(),
                            "tables": [],
                        })
            except Exception as e:
                logger.warning(f"PyPDF2 failed: {e}")

        # ── Phase 2: Tables via pdfplumber (per page) ──────────────────
        if PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page_idx, page in enumerate(pdf.pages):
                        page_num = page_idx + 1

                        # Ensure we have a matching pages_data entry
                        while len(pages_data) < page_num:
                            pages_data.append({"page": page_num, "text": "", "tables": []})

                        page_entry = pages_data[page_idx]

                        # Extract tables for this page
                        page_tables = page.extract_tables() or []
                        for raw_table in page_tables:
                            if not raw_table:
                                continue
                            # Clean cells (None → "")
                            clean_table = [
                                [str(cell).strip() if cell else "" for cell in row]
                                for row in raw_table
                            ]
                            # De-duplicate tables by content hash
                            table_hash = hashlib.md5(str(clean_table).encode()).hexdigest()
                            if table_hash in seen_table_hashes:
                                continue
                            seen_table_hashes.add(table_hash)

                            page_entry["tables"].append(clean_table)
                            all_tables.append(clean_table)

                        # If pdfplumber text is better (more chars), use it
                        plumber_text = page.extract_text() or ""
                        if len(plumber_text) > len(page_entry.get("text", "")):
                            page_entry["text"] = plumber_text.strip()

                        # Append a compact text version of the tables to this page's text
                        # so the LLM sees tables in context.
                        for table in page_entry["tables"]:
                            table_text = "\n".join(" | ".join(row) for row in table if any(c.strip() for c in row))
                            if table_text:
                                page_entry["text"] += "\n\n[TABLE]\n" + table_text

            except Exception as e:
                logger.warning(f"pdfplumber table extraction failed: {e}")

        # Fallback: pdfplumber text only (if no pages extracted at all)
        if not pages_data and PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page_idx, page in enumerate(pdf.pages):
                        text = page.extract_text() or ""
                        pages_data.append({
                            "page": page_idx + 1,
                            "text": text.strip(),
                            "tables": [],
                        })
            except Exception as e:
                logger.warning(f"pdfplumber text fallback failed: {e}")

        # ── Phase 3: Assemble full text ────────────────────────────────
        full_text = "\n\n".join(p["text"] for p in pages_data if p["text"])
        word_count = len(full_text.split())

        # ── Phase 4: OCR Fallback for scanned PDFs ─────────────────────
        if word_count < MIN_WORDS_FOR_TEXT_PDF:
            logger.info(
                f"Only {word_count} words extracted — likely a scanned PDF. "
                "Triggering Gemini OCR fallback..."
            )
            ocr_text = self._gemini_ocr_fallback(file_path)
            if ocr_text:
                full_text = ocr_text
                word_count = len(full_text.split())
                # Replace pages_data with OCR text as a single page
                pages_data = [{"page": 1, "text": full_text, "tables": []}]

        return {
            "extracted_text": full_text,
            "page_count": len(pages_data),
            "word_count": word_count,
            "pages": pages_data,
            "tables": all_tables,
            "processing_error": None,
        }

    def _gemini_ocr_fallback(self, file_path: str) -> str:
        """
        Use Gemini 1.5 Flash to perform OCR on a scanned PDF.
        """
        try:
            import google.generativeai as genai

            api_key = os.environ.get("GOOGLE_API_KEY")
            if not api_key or api_key == "missing_key_please_add_to_env":
                logger.warning("No GOOGLE_API_KEY found for OCR fallback.")
                return ""

            genai.configure(api_key=api_key)

            print("📤 Uploading document to Gemini for OCR...")
            uploaded_file = genai.upload_file(path=file_path, mime_type="application/pdf")

            model = genai.GenerativeModel("gemini-2.5-flash")

            print("🧠 Running OCR extraction...")
            response = model.generate_content([
                uploaded_file,
                (
                    "Extract all the text from this document accurately. "
                    "Do not summarize, just transcribe the text exactly as it appears. "
                    "Preserve the layout flow and table structures where possible."
                )
            ])

            try:
                genai.delete_file(uploaded_file.name)
            except Exception:
                pass

            return response.text

        except ImportError:
            logger.warning("google-generativeai not installed for OCR fallback.")
            return ""
        except Exception as e:
            logger.error(f"Gemini OCR fallback failed: {e}")
            return ""

    # ------------------------------------------------------------------
    # DOCX processing
    # ------------------------------------------------------------------

    def _process_docx(self, file_path: str) -> Dict:
        """
        Process DOCX file and extract text content and metadata.

        Improvements over previous version:
        - Preserves empty merged-cell slots so column count stays consistent.
        - De-duplicates tables by content hash.
        - Returns per-page approximation (by paragraph density).
        """
        if not DOCX_AVAILABLE:
            raise DocumentProcessingError("DOCX processing library not available")

        try:
            doc = DocxDocument(file_path)
            text_parts: List[str] = []
            tables_data: List[List[List[str]]] = []
            seen_table_hashes = set()
            
            # Extract images and process OCR
            extracted_photo_path = None
            try:
                images = self._extract_images_from_docx(file_path)
                if images:
                    # Find the most likely portrait photo (usually the largest portrait-shaped image)
                    best_image = images[0]
                    best_score = -1
                    try:
                        from PIL import Image
                        for img_path in images:
                            with Image.open(img_path) as img:
                                width, height = img.size
                                if width < 50 or height < 50:
                                    continue
                                
                                ratio = height / width if width else 0
                                area = width * height
                                
                                # Penalize wide/landscape images (like banners or logos) and extremely tall images
                                if ratio < 0.8:
                                    score = area * 0.01
                                elif ratio > 2.0:
                                    score = area * 0.01
                                else:
                                    # Perfect portrait is ~1.33. Score based on distance from perfect portrait.
                                    distance = abs(ratio - 1.33)
                                    multiplier = 100.0 / (1.0 + distance * 10)
                                    score = area * multiplier
                                    
                                if score > best_score:
                                    best_score = score
                                    best_image = img_path
                    except ImportError:
                        logger.warning("Pillow not installed. Falling back to first image.")
                    except Exception as e:
                        logger.warning(f"Error scoring images: {e}")
                        
                    extracted_photo_path = best_image
                    
                    # Process remaining images with Gemini OCR
                    if len(images) > 1:
                        logger.info(f"Found {len(images)-1} additional images in DOCX. Running OCR...")
                        for img_path in images:
                            if img_path == extracted_photo_path:
                                continue
                            ocr_text = self._gemini_ocr_image(img_path)
                            if ocr_text:
                                text_parts.append(ocr_text)
            except Exception as e:
                logger.warning(f"Failed to process images from DOCX: {e}")

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract tables
            for table in doc.tables:
                table_matrix: List[List[str]] = []

                for row in table.rows:
                    # Build a row with the EXACT number of columns (preserve empty merged cells)
                    raw_cells = list(row.cells)
                    clean_row: List[str] = []
                    prev_text = None

                    for cell in raw_cells:
                        txt = cell.text.strip()
                        # python-docx repeats merged cells — deduplicate sequentially
                        # BUT keep empty strings to preserve column count.
                        if txt == prev_text and txt != "":
                            # This is a merged cell continuation — insert empty placeholder
                            clean_row.append("")
                        else:
                            clean_row.append(txt)
                            prev_text = txt

                    if any(c.strip() for c in clean_row):
                        table_matrix.append(clean_row)
                        # Also append to running text
                        text_parts.append(" | ".join(c for c in clean_row if c.strip()))

                if table_matrix:
                    table_hash = hashlib.md5(str(table_matrix).encode()).hexdigest()
                    if table_hash not in seen_table_hashes:
                        seen_table_hashes.add(table_hash)
                        tables_data.append(table_matrix)

            full_text = "\n\n".join(text_parts)
            word_count = len(full_text.split())
            estimated_pages = max(1, round(word_count / 250))

            return {
                "extracted_text": full_text,
                "page_count": estimated_pages,
                "word_count": word_count,
                # Simulate per-page data as a single entry for DOCX
                "pages": [{"page": 1, "text": full_text, "tables": tables_data}],
                "tables": tables_data,
                "processing_error": None,
                "extracted_photo_path": extracted_photo_path,
            }
        except Exception as e:
            logger.error(f"Error processing DOCX: {e}")
            raise DocumentProcessingError(f"Failed to process DOCX file: {str(e)}")

    def _extract_images_from_docx(self, file_path: str) -> List[str]:
        """Extracts images from a DOCX file and saves them temporarily."""
        images = []
        try:
            with zipfile.ZipFile(file_path, 'r') as docx_zip:
                for item in docx_zip.namelist():
                    if item.startswith('word/media/') and item.lower().endswith(('.png', '.jpg', '.jpeg')):
                        # Extract the image
                        temp_dir = tempfile.mkdtemp(prefix="docx_images_")
                        extracted_path = docx_zip.extract(item, temp_dir)
                        # The extracted path is inside temp_dir/word/media/...
                        # We just return the full path to the extracted file
                        images.append(extracted_path)
        except zipfile.BadZipFile:
            logger.warning(f"Bad zip file for DOCX image extraction: {file_path}")
        return images
        
    def _gemini_ocr_image(self, image_path: str) -> str:
        """Run OCR on a single image extracted from DOCX using Gemini."""
        try:
            import google.generativeai as genai
            api_key = os.environ.get("GOOGLE_API_KEY")
            if not api_key or api_key == "missing_key_please_add_to_env":
                return ""
                
            genai.configure(api_key=api_key)
            uploaded_file = genai.upload_file(path=image_path)
            print("📤 Uploading document to Gemini for OCR...")
            model = genai.GenerativeModel("gemini-2.5-flash")
            response = model.generate_content([
                uploaded_file,
                "Extract all text and structured data from this image accurately. Do not summarize."
            ])
            try:
                genai.delete_file(uploaded_file.name)
            except Exception:
                pass
            return response.text
        except Exception as e:
            logger.error(f"Gemini OCR on image failed: {e}")
            return ""
    # ------------------------------------------------------------------
    # Utility
    # ------------------------------------------------------------------

    def get_document_info(self, file_path: str) -> Dict:
        """Get basic information about a document without full processing."""
        if not os.path.exists(file_path):
            raise DocumentProcessingError(f"File not found: {file_path}")

        file_stat = os.stat(file_path)
        document_type = self._detect_document_type(file_path)

        return {
            'file_size': file_stat.st_size,
            'document_type': document_type,
            'file_extension': Path(file_path).suffix.lower(),
            'file_name': Path(file_path).name,
            'supported': document_type in self.supported_formats,
        }


class AsyncDocumentProcessor:
    """
    Asynchronous wrapper for document processing.
    Useful for processing large documents without blocking the main thread.
    """

    def __init__(self):
        self.processor = DocumentProcessor()

    async def process_document_async(self, file_path: str, document_type: str = None) -> Dict:
        import asyncio
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None,
            self.processor.process_document,
            file_path,
            document_type,
        )


# Example usage and testing
if __name__ == "__main__":
    processor = DocumentProcessor()
    print("Supported formats:", processor.supported_formats)
