import os
import sys
import json
import re
from datetime import datetime

# Add the project root to the sys.path
project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.append(project_root)

# Set the Django settings module
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'saker.settings')

import django
# Override logging so it doesn't try to access production logs when running locally
from django.conf import settings
settings.LOGGING_CONFIG = None
django.setup()

from api.models import Users, CVSubmission, Rank, UserRank, SeaService, Certificate
from api.serializer import CVSubmissionSerializer

def clean_text(text):
    if text is None: return ""
    return str(text).strip().replace('\n', ' ')

def safe_date(date_str):
    if not date_str:
        return None
    date_str = clean_text(date_str).lower()
    if date_str in ["none", "null", "n/a", "unlimited", "until now", "present", ""]:
        return None
    if "x" in date_str or "×" in date_str or "*" in date_str:
        return None
        
    # Try multiple formats
    patterns = [
        r'^(\d{2})[-/](\d{2})[-/](\d{4})$',  # DD-MM-YYYY
        r'^(\d{4})[-/](\d{2})[-/](\d{2})$',  # YYYY-MM-DD
    ]
    for p in patterns:
        match = re.match(p, date_str)
        if match:
            # Reformat to YYYY-MM-DD
            if len(match.group(1)) == 4:
                return f"{match.group(1)}-{match.group(2)}-{match.group(3)}"
            else:
                return f"{match.group(3)}-{match.group(2)}-{match.group(1)}"
    return None

def extract_from_docx(file_path):
    import docx
    try:
        doc = docx.Document(file_path)
    except Exception as e:
        print(f"Error reading docx: {e}")
        return None

    data = {
        "user_name": "",
        "user_email_display": "",
        "position_name": "Other",
        "status": "Pending",
        "notes": "Auto-submitted via local DOCX parser",
        "seafarer_application": {
            "1_personal_details": {"full_name": "", "date_of_birth": None, "nationality": "", "place_of_birth": "", "marital_status": {"single": False, "married": False}},
            "3_contact_details": {"home_address_city": "", "e_mail": "", "mobile_tel": ""},
            "4_travel_documents": [],
            "5_professional_qualification_certificate_of_competency": [],
            "8_marine_courses": [],
            "9_complete_sea_service_details": {"applicant_info": {}, "service_records": []}
        },
        "user_documents": {
            "passport": {}, "seaman_book": {}, "coc": {}, "goc": {}, "health_certificate": {}, "sea_service": [], "marine_courses": []
        }
    }

    # Helper to deduplicate row cells (due to merged cells in word)
    def get_row_cells(row):
        cells = []
        for cell in row.cells:
            t = clean_text(cell.text)
            if not cells or cells[-1] != t:
                cells.append(t)
        return cells

    for table in doc.tables:
        text_content = " ".join([clean_text(c.text) for r in table.rows for c in r.cells])
        
        # 1. Position and Register Code
        if "Application For Position" in text_content:
            for row in table.rows:
                cells = get_row_cells(row)
                if len(cells) >= 2:
                    if "Application For Position" in cells[0]:
                        data["position_name"] = cells[1]
                    if "Register Code" in cells[0]:
                        data["register_code"] = cells[1]

        # 2. Personal Details
        if "PERSONAL DETAILS" in text_content and "Full Name" in text_content:
            for row in table.rows:
                cells = get_row_cells(row)
                for i, c in enumerate(cells):
                    if c.startswith("Full Name") and i+1 < len(cells) and not data["user_name"]:
                        data["user_name"] = cells[i+1]
                        data["seafarer_application"]["1_personal_details"]["full_name"] = cells[i+1]
                    if c.startswith("Date Of Birth") and i+1 < len(cells):
                        data["seafarer_application"]["1_personal_details"]["date_of_birth"] = safe_date(cells[i+1])
                    if "Single" in c and "✓" in c:
                        data["seafarer_application"]["1_personal_details"]["marital_status"]["single"] = True
                    if "Married" in c and "✓" in c:
                        data["seafarer_application"]["1_personal_details"]["marital_status"]["married"] = True
                    if "Nationality" in c and i+1 < len(cells):
                        data["seafarer_application"]["1_personal_details"]["nationality"] = cells[i+1]
        
        # 3. Contact Details
        if "CONTACT DETAILS" in text_content:
            # We want to avoid matching "Next of Kin" mobile, which appears later in the same table.
            # So we only set the mobile_tel if it's not already set, or if we specifically match "Mobile / Tel"
            for row in table.rows:
                cells = get_row_cells(row)
                for i, c in enumerate(cells):
                    if "e-mail" in c.lower() and i+1 < len(cells) and "@" in cells[i+1]:
                        data["user_email_display"] = cells[i+1]
                        data["seafarer_application"]["3_contact_details"]["e_mail"] = cells[i+1]
                    elif "mobile" in c.lower() and i+1 < len(cells):
                        # only set if it hasn't been set yet, since candidate mobile is above Next of Kin
                        if not data["seafarer_application"]["3_contact_details"].get("mobile_tel"):
                            data["seafarer_application"]["3_contact_details"]["mobile_tel"] = cells[i+1]

        # 4. Travel Documents
        if "TRAVEL DOCUMENTS" in text_content and "Passport" in text_content:
            for row in table.rows:
                cells = get_row_cells(row)
                if len(cells) >= 4:
                    doc_type = cells[0].lower()
                    if "passport" in doc_type:
                        data["user_documents"]["passport"] = {
                            "passport_no": cells[1], "issue_date": safe_date(cells[2]), "expiry_date": safe_date(cells[3])
                        }
                    elif "seaman" in doc_type:
                        data["user_documents"]["seaman_book"] = {
                            "seaman_book_no": cells[1], "issue_date": safe_date(cells[2]), "expiry_date": safe_date(cells[3])
                        }

        # 5. COC / GOC
        if "CERTIFICATE OF COMPETENCY" in text_content:
            for row in table.rows:
                cells = get_row_cells(row)
                if len(cells) >= 4:
                    if "COC" in cells[0].upper():
                        data["user_documents"]["coc"] = {
                            "certificate_name": cells[0], "certificate_number": cells[1], 
                            "issue_date": safe_date(cells[2]), "expiry_date": safe_date(cells[3])
                        }
                    elif "GOC" in cells[0].upper():
                        data["user_documents"]["goc"] = {
                            "certificate_number": cells[1], "issue_date": safe_date(cells[2]), "expiry_date": safe_date(cells[3])
                        }

        # 8. Marine Courses
        if "MARINE COURSES" in text_content:
            for row in table.rows:
                cells = get_row_cells(row)
                if len(cells) >= 3 and cells[0] and cells[0] != "Course Name":
                    if "EAMS / Alex" not in cells[0]: # filter out noisy headers
                        data["user_documents"]["marine_courses"].append({
                            "course_name": cells[0],
                            "number": cells[1] if len(cells)>1 else "",
                            "issue_date": safe_date(cells[2]) if len(cells)>2 else None,
                            "expiry_date": safe_date(cells[3]) if len(cells)>3 else None
                        })

        # 9. Sea Service
        if "COMPLETE SEA" in text_content:
            for row in table.rows:
                cells = get_row_cells(row)
                if len(cells) >= 5 and cells[0] and "Company Name" not in cells[0]:
                    # Ignore rows that look like contact info (references)
                    if "@" in cells[4] or re.search(r'\d{10}', cells[3]):
                        continue
                        
                    data["user_documents"]["sea_service"].append({
                        "company_name": cells[0],
                        "rank": cells[1] if len(cells)>1 else "",
                        "vessel_name": cells[2] if len(cells)>2 else "",
                        "signed_on": safe_date(cells[3]) if len(cells)>3 else None,
                        "signed_off": safe_date(cells[4]) if len(cells)>4 else None,
                        "vessel_type": cells[5] if len(cells)>5 else "",
                        "dwt": cells[6] if len(cells)>6 else ""
                    })

    # Robust Regex fallback for email: scan ALL text in the document
    if not data["user_email_display"]:
        full_text = []
        for p in doc.paragraphs:
            if p.text.strip(): full_text.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip(): full_text.append(cell.text)
        
        combined_text = " ".join(full_text)
        match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', combined_text)
        if match:
            extracted_email = match.group(0).lower()
            data["user_email_display"] = extracted_email
            data["seafarer_application"]["3_contact_details"]["e_mail"] = extracted_email

    return data

def run_ingestion(ai_data, file_name=""):
    email_to_check = ai_data.get("user_email_display") or ""
    import re
    is_valid_email = bool(re.match(r"^[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+$", str(email_to_check)))
    
    if not email_to_check or not is_valid_email:
        print(f"[{file_name}] Error: No email provided in data or email is invalid.")
        return False

    email = email_to_check

    user_name = ai_data.get("user_name") or "Applicant"
    parts = user_name.split(" ", 1)
    user_first_name = parts[0] if parts[0].strip() else "Applicant"
    user_middle_name = parts[1] if len(parts) > 1 and parts[1].strip() else "Unknown"

    user, created = Users.objects.get_or_create(
        email=email,
        defaults={
            "first_name": user_first_name, 
            "role": "Employee",
            "coc_issued_by": ""
        }
    )
    
    app_data = ai_data.get("seafarer_application", {})
    personal_details = app_data.get("1_personal_details", {})
    contact_details = app_data.get("3_contact_details", {})

    if personal_details.get("date_of_birth"): user.date_of_birth = personal_details.get("date_of_birth")
    if personal_details.get("nationality"): user.nationality = personal_details.get("nationality")
    if personal_details.get("place_of_birth"): user.Place_Of_Birth = personal_details.get("place_of_birth")
    
    marital = personal_details.get("marital_status", {})
    if marital.get("married"): user.marital_status = "Married"
    elif marital.get("single"): user.marital_status = "Single"

    if contact_details.get("home_address_city"): user.city = contact_details.get("home_address_city")
    if contact_details.get("mobile_tel"): user.phone_number = contact_details.get("mobile_tel")

    if not getattr(user, 'coc_issued_by', "fallback"):
        user.coc_issued_by = ""

    user.save()

    def clean_doc(d):
        if not isinstance(d, dict):
            return d
        return {k: v if v is not None else "" for k, v in d.items()}

    payload = {
        "user_first_name": user_first_name,
        "user_middle_name": user_middle_name,
        "user_email": email,
        "position": ai_data.get("position_name"),
        "status": ai_data.get("status") or "Pending",
        "experience_years": ai_data.get("experience_years", 0),
    }

    user_docs = ai_data.get("user_documents", {})
    if clean_doc(user_docs.get("passport")): payload["passport_update"] = clean_doc(user_docs.get("passport"))
    if clean_doc(user_docs.get("seaman_book")): payload["seaman_book_update"] = clean_doc(user_docs.get("seaman_book"))
    if clean_doc(user_docs.get("other_seaman_book")): payload["other_seaman_book_update"] = clean_doc(user_docs.get("other_seaman_book"))
    if clean_doc(user_docs.get("coc")): payload["coc_update"] = clean_doc(user_docs.get("coc"))
    if clean_doc(user_docs.get("goc")): payload["goc_update"] = clean_doc(user_docs.get("goc"))

    for key in ['passport_update', 'seaman_book_update', 'coc_update', 'goc_update', 'other_seaman_book_update']:
        if payload.get(key):
            for subkey in ['issue_date', 'expiry_date']:
                if subkey in payload[key] and not payload[key][subkey]:
                    payload[key][subkey] = None
        else:
            payload.pop(key, None)

    position_name = payload.get("position")
    rank_obj = None
    if position_name:
        rank_obj = Rank.objects.filter(name__iexact=position_name).first()
        if not rank_obj:
            import uuid
            for _ in range(10):
                code = f"CUS-{str(uuid.uuid4())[:6].upper()}"
                if not Rank.objects.filter(code=code).exists():
                    rank_obj = Rank.objects.create(name=position_name, code=code)
                    break

    if rank_obj:
        existing_submission = CVSubmission.objects.filter(user=user, position=rank_obj).first()
    else:
        existing_submission = CVSubmission.objects.filter(user=user).first()

    if existing_submission:
        serializer = CVSubmissionSerializer(instance=existing_submission, data=payload, partial=True)
    else:
        serializer = CVSubmissionSerializer(data=payload)

    if serializer.is_valid():
        submission = serializer.save(user=user)
        print(f"[{file_name}] Successfully ingested CV Submission ID: {submission.id}")
    else:
        print(f"[{file_name}] Serializer validation failed:", json.dumps(serializer.errors, indent=2))
        return False

    sea_service_list = ai_data.get("user_documents", {}).get("sea_service", [])
    if sea_service_list and isinstance(sea_service_list, list):
        from api.models import SeaService
        from core.models import VesselType
        from core.models import Flag
        sea_service_count = 0
        for ss in sea_service_list:
            if not isinstance(ss, dict):
                continue
            v_type_str = ss.get("vessel_type")
            if not v_type_str:
                v_type_str = "Unknown"
            v_type_obj, _ = VesselType.objects.get_or_create(name=v_type_str)
            
            flag_str = ss.get("flag")
            if not flag_str:
                flag_str = "Unknown"
            flag_obj, _ = Flag.objects.get_or_create(name=flag_str)
                
            rank_str = ss.get("rank")
            if not rank_str:
                rank_str = "Unknown"
            rank_obj = Rank.objects.filter(name__iexact=rank_str).first()
            if not rank_obj:
                import uuid
                rank_obj = Rank.objects.create(code=f"UNK-{str(uuid.uuid4())[:4].upper()}", name=rank_str)
            
            try:
                SeaService.objects.get_or_create(
                    user=user,
                    vessel_name=ss.get("vessel_name") or "Unknown",
                    rank=rank_obj,
                    vessel_type=v_type_obj,
                    flag=flag_obj,
                    defaults={
                        "signed_on": ss.get("signed_on") if ss.get("signed_on") else None,
                        "signed_off": ss.get("signed_off") if ss.get("signed_off") else None,
                    }
                )
                sea_service_count += 1
            except Exception as e:
                print(f"[{file_name}] Warning: Could not save sea service due to {e}")
        print(f"Ingested {sea_service_count} new Sea Service records.")

    marine_courses_list = ai_data.get("user_documents", {}).get("marine_courses", [])
    if marine_courses_list and isinstance(marine_courses_list, list):
        try:
            from courses.models import Course
            marine_course_count = 0
            for c in marine_courses_list:
                if not isinstance(c, dict):
                    continue
                course_name = c.get("course_name")
                if course_name:
                    try:
                        Course.objects.get_or_create(
                            user=user,
                            course_name=course_name,
                            defaults={
                                "issue_date": c.get("issue_date") if c.get("issue_date") else None,
                                "expiry_date": c.get("expiry_date") if c.get("expiry_date") else None,
                            }
                        )
                        marine_course_count += 1
                    except Exception as e:
                        print(f"[{file_name}] Warning: Could not save marine course due to {e}")
            print(f"Ingested {marine_course_count} new Marine Course records.")
        except ImportError:
            pass

    return True

def process_folder(folder_path="all_json_files"):
    import shutil
    folder_path = os.path.join(project_root, folder_path)
    processed_folder = os.path.join(project_root, f"{os.path.basename(folder_path)}_processed")
    failed_folder = os.path.join(project_root, f"{os.path.basename(folder_path)}_failed")
    os.makedirs(processed_folder, exist_ok=True)
    os.makedirs(failed_folder, exist_ok=True)
    
    files = [f for f in os.listdir(folder_path) if f.endswith(('.json', '.docx'))]
    print(f"Found {len(files)} files in {folder_path}. Starting processing...\n")

    success_count = fail_count = 0
    failed_files_list = []

    for filename in files:
        file_path = os.path.join(folder_path, filename)
        file_ext = os.path.splitext(filename)[1].lower()
        print(f"--- Processing {filename} ---")
        try:
            if file_ext == '.json':
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
            else:
                data = extract_from_docx(file_path)
            
            if not data:
                print(f"[{filename}] Error: Data extraction returned None.")
                fail_count += 1
                failed_files_list.append(f"{filename} (Extraction returned None)")
                shutil.move(file_path, os.path.join(failed_folder, filename))
                continue
                
            if run_ingestion(data, file_name=filename):
                shutil.move(file_path, os.path.join(processed_folder, filename))
                success_count += 1
            else:
                fail_count += 1
                failed_files_list.append(f"{filename} (Ingestion failed)")
                shutil.move(file_path, os.path.join(failed_folder, filename))
        except Exception as e:
            print(f"[{filename}] Error: {e}")
            fail_count += 1
            failed_files_list.append(f"{filename} (Exception: {e})")
            shutil.move(file_path, os.path.join(failed_folder, filename))

    print(f"\nBatch complete! Success: {success_count}, Failed: {fail_count}")
    
    if failed_files_list:
        failed_log_path = os.path.join(project_root, "failed_files.txt")
        with open(failed_log_path, "w", encoding="utf-8") as f:
            f.write(f"Batch failed files ({datetime.now().strftime('%Y-%m-%d %H:%M:%S')}):\n")
            for failed_file in failed_files_list:
                f.write(f"- {failed_file}\n")
        print(f"\nA list of all failed files has been saved to: {failed_log_path}")

def cleanup_duplicates():
    from django.db.models import Count
    # Find users with multiple CV Submissions for the same position
    duplicates = CVSubmission.objects.values('user', 'position').annotate(
        count=Count('id')
    ).filter(count__gt=1)

    for dup in duplicates:
        user_id = dup['user']
        pos_id = dup['position']
        # Get all submissions for this user/position ordered by latest first
        submissions = CVSubmission.objects.filter(
            user_id=user_id, position_id=pos_id
        ).order_by('-id')
        
        # Keep the first one (most recent), delete the rest
        if submissions.count() > 1:
            to_keep = submissions.first()
            to_delete = submissions.exclude(id=to_keep.id)
            deleted_count, _ = to_delete.delete()
            print(f"Cleaned up {deleted_count} duplicate CV Submissions for User ID {user_id}.")

if __name__ == "__main__":
    folder = sys.argv[1] if len(sys.argv) > 1 else "all_json_files"
    process_folder(folder)
    print("Running duplicate cleanup just in case...")
    cleanup_duplicates()
